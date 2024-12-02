import openai
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

class PRFAQGenerator:
    def __init__(self, api_key):
        openai.api_key = api_key

    def _generate_section(self, prompt, product_desc):
        try:
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert at writing Amazon-style PRFAQ documents. "
                                                "Provide detailed, professional content following Amazon's standards."},
                    {"role": "user", "content": f"{prompt}\n\nProduct Description:\n{product_desc}"}
                ],
                temperature=0.7,
                max_tokens=2000
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            return f"Error generating content: {str(e)}"

    def generate_press_release(self, product_desc):
        prompt = """Write a press release for this product following Amazon's press release format. 
        Include a compelling headline, location/date, opening paragraph, problem statement, 
        solution description, customer quote, and company quote."""
        return self._generate_section(prompt, product_desc)

    def generate_faq(self, product_desc):
        prompt = """Generate a comprehensive FAQ section for this product. Include at least 8 questions 
        covering technical aspects, user benefits, implementation details, and potential concerns."""
        return self._generate_section(prompt, product_desc)

    def generate_working_backwards(self, product_desc):
        prompt = """Create a Working Backwards section that explains the customer-centric approach taken 
        in developing this product. Include customer personas, key use cases, and success criteria."""
        return self._generate_section(prompt, product_desc)

    def generate_technical_architecture(self, product_desc):
        prompt = """Describe the technical architecture for this product. Include system components, 
        data flow, security considerations, and scalability aspects."""
        return self._generate_section(prompt, product_desc)

    def generate_metrics(self, product_desc):
        prompt = """Define key success metrics for this product. Include both business and technical KPIs, 
        target values, and measurement methods."""
        return self._generate_section(prompt, product_desc)

    def generate_go_to_market(self, product_desc):
        prompt = """Create a go-to-market strategy section. Include target market analysis, 
        marketing channels, launch timeline, and competitive positioning."""
        return self._generate_section(prompt, product_desc)

    def create_document(self, product_desc):
        # Generate all sections
        sections = {
            "Press Release": self.generate_press_release(product_desc),
            "FAQ": self.generate_faq(product_desc),
            "Working Backwards": self.generate_working_backwards(product_desc),
            "Technical Architecture": self.generate_technical_architecture(product_desc),
            "Success Metrics": self.generate_metrics(product_desc),
            "Go-to-Market Strategy": self.generate_go_to_market(product_desc)
        }
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('PRFAQ Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_paragraph.add_run(datetime.now().strftime("%B %d, %Y"))
        
        # Add sections
        for section_title, content in sections.items():
            doc.add_heading(section_title, level=1)
            doc.add_paragraph(content)
            doc.add_paragraph()  # Add spacing between sections
        
        return doc

    def save_document(self, doc, filename):
        doc.save(filename)
        return filename
